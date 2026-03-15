"""
Profile 文档解析：PDF、DOCX、TXT、MD、XLSX、CSV。
提取文本后通过规则映射为结构化字段 + spec_items。
"""
import io
import re
from typing import Any

import pandas as pd

# 可选依赖（Profile 上传 PDF/DOCX 需要）
_pdfplumber_error: str | None = None
try:
    import pdfplumber
except ImportError as e:
    pdfplumber = None
    _pdfplumber_error = str(e)
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

# 中英文参数名映射到统一 key
FIELD_ALIASES = {
    "brand": ["brand", "品牌"],
    "model": ["model", "型号", "model number"],
    "full_name": ["full_name", "产品名称", "name", "机型", "手机型号"],
    "display_type": ["display_type", "display", "屏幕类型", "屏幕", "screen type"],
    "display_size": ["display_size", "屏幕尺寸", "size", "screen size"],
    "resolution": ["resolution", "分辨率", "分辨率"],
    "refresh_rate": ["refresh_rate", "刷新率", "refresh rate"],
    "os": ["os", "操作系统", "操作系统", "android", "ios"],
    "chipset": ["chipset", "soc", "平台", "芯片", "处理器平台"],
    "cpu": ["cpu", "处理器", "processor"],
    "gpu": ["gpu", "图形"],
    "battery": ["battery", "电池", "电池容量", "battery capacity"],
    "charging": ["charging", "充电", "快充", "charging"],
    "main_camera": ["main_camera", "rear camera", "后摄", "主摄", "主摄像头", "main camera"],
    "selfie_camera": ["selfie_camera", "front camera", "前摄", "自拍", "前置"],
    "memory_summary": ["memory_summary", "memory", "存储", "内存", "ram", "rom"],
    "price": ["price", "价格", "参考价"],
    "launch_date": ["launch_date", "launch", "发布时间", "上市"],
    "weight": ["weight", "重量"],
    "dimensions": ["dimensions", "尺寸"],
}


def _normalize_field_key(key: str) -> str | None:
    """将用户/文档中的参数名映射为统一字段。"""
    k = (key or "").strip().lower()
    for canonical, aliases in FIELD_ALIASES.items():
        for a in aliases:
            if a.lower() in k or k in a.lower():
                return canonical
    return None


def extract_text_from_file(content: bytes, filename: str) -> tuple[bool, str, str | None]:
    """
    根据扩展名解析文件为纯文本。
    返回: (success, text, error_message)
    """
    fn = (filename or "").lower()
    if fn.endswith(".txt") or fn.endswith(".md"):
        try:
            return True, content.decode("utf-8", errors="replace"), None
        except Exception as e:
            return False, "", str(e)
    if fn.endswith(".docx"):
        if DocxDocument is None:
            return False, "", "请安装 python-docx: pip install python-docx"
        try:
            doc = DocxDocument(io.BytesIO(content))
            text = "\n".join(p.text for p in doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    text += "\n" + "\t".join(cell.text for cell in row.cells)
            return True, text, None
        except Exception as e:
            return False, "", f"DOCX 解析失败: {e}"
    if fn.endswith(".pdf"):
        if pdfplumber is None:
            hint = "请在本项目虚拟环境中安装: pip install pdfplumber"
            if _pdfplumber_error:
                hint += f"（当前错误: {_pdfplumber_error}）"
            return False, "", hint
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                parts = []
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        parts.append(t)
                return True, "\n".join(parts), None
        except Exception as e:
            return False, "", f"PDF 解析失败: {e}"
    if fn.endswith(".xlsx") or fn.endswith(".xls"):
        try:
            df = pd.read_excel(io.BytesIO(content), header=None)
            return True, df.to_string(), None
        except Exception as e:
            return False, "", f"Excel 解析失败: {e}"
    if fn.endswith(".csv"):
        try:
            df = pd.read_csv(io.BytesIO(content), encoding="utf-8", on_bad_lines="skip", header=None)
            return True, df.to_string(), None
        except Exception:
            try:
                df = pd.read_csv(io.BytesIO(content), encoding="gbk", on_bad_lines="skip", header=None)
                return True, df.to_string(), None
            except Exception as e:
                return False, "", f"CSV 解析失败: {e}"
    return False, "", "不支持的文件格式，请上传 PDF / DOCX / TXT / MD / XLSX / CSV"


def parse_text_to_specs(text: str) -> tuple[dict[str, str | None], list[dict]]:
    """
    从纯文本中按行或键值对提取字段。
    - 返回 (fields_dict, spec_items_list)
    - 规则：行内包含 :、：、\t 的拆分为 key: value；常见参数名做映射。
    """
    fields = {}
    spec_items = []
    seen = set()
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        for sep in (":", "：", "\t"):
            if sep in line:
                parts = line.split(sep, 1)
                if len(parts) == 2:
                    key = parts[0].strip()
                    val = parts[1].strip()
                    canonical = _normalize_field_key(key)
                    if canonical and canonical not in seen:
                        seen.add(canonical)
                        fields[canonical] = val or None
                    spec_items.append({"spec_group": "Imported", "spec_key": key, "spec_value": val, "sort_order": len(spec_items)})
                break
        else:
            # 整行当作值，尝试匹配常见键
            for canonical, aliases in FIELD_ALIASES.items():
                for a in aliases:
                    if a.lower() in line[:30]:
                        if canonical not in seen:
                            seen.add(canonical)
                            fields[canonical] = line
                        spec_items.append({"spec_group": "Imported", "spec_key": canonical, "spec_value": line, "sort_order": len(spec_items)})
                        break
    return fields, spec_items


def parse_profile_file(
    content: bytes, filename: str, use_ai: bool = False
) -> tuple[bool, dict[str, Any], str | None]:
    """
    上传文件后调用：先提取文本，再解析为 fields + spec_items。
    use_ai=True 时调用 AI 解析（需配置 DeepSeek API），适配不同 Profile 格式；
    否则使用规则解析。AI 失败时自动回退到规则解析。
    返回: (success, { "file_name", "fields", "spec_items", "message", "parse_mode" }, error)
    """
    ok, text, err = extract_text_from_file(content, filename)
    if not ok:
        return False, {"file_name": filename, "fields": {}, "spec_items": [], "message": err}, err

    text = (text or "").strip()
    if len(text) < 30:
        return False, {
            "file_name": filename,
            "fields": {},
            "spec_items": [],
            "message": "文档中未提取到足够文字（可能为扫描版 PDF 或空白页），请上传可复制文字的 PDF 或使用 DOCX/TXT。",
        }, "文档中未提取到足够文字"

    if use_ai:
        try:
            from app.services import ai_service
            ai_ok, ai_fields, ai_spec_items, ai_err = ai_service.parse_profile_with_ai(text)
            if ai_ok:
                return True, {
                    "file_name": filename,
                    "fields": ai_fields,
                    "spec_items": ai_spec_items,
                    "message": "AI 解析完成，请确认后保存",
                    "parse_mode": "ai",
                }, None
            # AI 失败则回退到规则解析，并在 message 中提示
            fields, spec_items = parse_text_to_specs(text)
            return True, {
                "file_name": filename,
                "fields": fields,
                "spec_items": spec_items,
                "message": f"AI 解析未可用（{ai_err}），已用规则解析。请确认后保存",
                "parse_mode": "rule_fallback",
            }, None
        except Exception as e:
            fields, spec_items = parse_text_to_specs(text)
            return True, {
                "file_name": filename,
                "fields": fields,
                "spec_items": spec_items,
                "message": f"AI 解析异常（{e}），已用规则解析。请确认后保存",
                "parse_mode": "rule_fallback",
            }, None

    fields, spec_items = parse_text_to_specs(text)
    return True, {
        "file_name": filename,
        "fields": fields,
        "spec_items": spec_items,
        "message": "解析完成，请确认后保存",
        "parse_mode": "rule",
    }, None
